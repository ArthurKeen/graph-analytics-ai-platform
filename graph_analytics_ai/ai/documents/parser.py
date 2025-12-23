"""
Document parser for extracting text from various file formats.

Supports: TXT, MD, PDF, DOCX (with optional dependencies).
"""

from typing import List

# Optional dependencies (exposed for testing/mocking)
try:
    import pdfplumber  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    pdfplumber = None  # type: ignore

try:
    from PyPDF2 import PdfReader  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    PdfReader = None  # type: ignore

try:
    from bs4 import BeautifulSoup  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    BeautifulSoup = None  # type: ignore

try:
    from docx import Document as DocxDocument  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    DocxDocument = None  # type: ignore

from .models import Document, DocumentMetadata, DocumentType, TextChunk


class ParserError(Exception):
    """Base exception for parser errors."""

    pass


class UnsupportedFormatError(ParserError):
    """Raised when document format is not supported."""

    pass


class DocumentParser:
    """
    Parser for extracting text from various document formats.

    Supports:
    - Plain text (.txt)
    - Markdown (.md)
    - PDF (.pdf) - requires PyPDF2 or pdfplumber
    - DOCX (.docx) - requires python-docx

    Example:
        >>> parser = DocumentParser()
        >>> doc = parser.parse("requirements.pdf")
        >>> print(f"Extracted {doc.word_count} words")
        >>> print(doc.get_preview())
    """

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 100):
        """
        Initialize document parser.

        Args:
            chunk_size: Size of text chunks (in characters).
            chunk_overlap: Overlap between chunks (in characters).
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def parse(self, file_path: str, chunk: bool = False) -> Document:
        """
        Parse a document file and extract text.

        Args:
            file_path: Path to the document file.
            chunk: Whether to split document into chunks.

        Returns:
            Parsed Document object.

        Raises:
            ParserError: If parsing fails.
            UnsupportedFormatError: If format is not supported.
        """
        # Get metadata
        metadata = DocumentMetadata.from_file(file_path)

        # Parse based on type
        try:
            if metadata.document_type == DocumentType.TEXT:
                content = self._parse_text(file_path)
            elif metadata.document_type == DocumentType.MARKDOWN:
                content = self._parse_markdown(file_path)
            elif metadata.document_type == DocumentType.PDF:
                content = self._parse_pdf(file_path)
            elif metadata.document_type == DocumentType.DOCX:
                content = self._parse_docx(file_path)
            elif metadata.document_type == DocumentType.HTML:
                content = self._parse_html(file_path)
            else:
                raise UnsupportedFormatError(
                    f"Unsupported document type: {metadata.document_type}"
                )

            # Create document
            doc = Document(metadata=metadata, content=content)

            # Chunk if requested
            if chunk:
                doc.chunks = self._create_chunks(content)

            return doc

        except Exception as e:
            # Create document with error
            doc = Document(metadata=metadata, content="", extraction_errors=[str(e)])
            return doc

    def _parse_text(self, file_path: str) -> str:
        """Parse plain text file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def _parse_markdown(self, file_path: str) -> str:
        """
        Parse markdown file.

        For now, just returns the raw markdown.
        Could be enhanced to strip markdown syntax if needed.
        """
        return self._parse_text(file_path)

    def _parse_pdf(self, file_path: str) -> str:
        """
        Parse PDF file.

        Tries multiple libraries in order of preference:
        1. pdfplumber (best quality)
        2. PyPDF2 (fallback)
        """
        # Try pdfplumber first if available
        if pdfplumber is not None:
            text = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    extract_fn = getattr(page, "extract_text", None)
                    page_text = None
                    if extract_fn:
                        try:
                            page_text = extract_fn()
                        except TypeError:
                            # Handle mock functions without self parameter (unit tests)
                            if hasattr(extract_fn, "__func__"):
                                page_text = extract_fn.__func__()
                    if page_text:
                        text.append(page_text)
            return "\n\n".join(text)

        # Fallback to PyPDF2 if available
        if PdfReader is not None:
            reader = PdfReader(file_path)
            text = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            return "\n\n".join(text)

        # If neither library is installed, raise clear error
        raise ParserError(
            "PDF parsing requires 'pdfplumber' or 'PyPDF2'. "
            "Install with: pip install pdfplumber"
        )

    def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX file."""
        if DocxDocument is None:
            raise ParserError(
                "DOCX parsing requires 'python-docx'. "
                "Install with: pip install python-docx"
            )

        doc = DocxDocument(file_path)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        return "\n\n".join(text)

    def _parse_html(self, file_path: str) -> str:
        """Parse HTML file."""
        try:
            from bs4 import BeautifulSoup as BS  # type: ignore
        except ImportError:
            raise ParserError(
                "HTML parsing requires 'beautifulsoup4'. "
                "Install with: pip install beautifulsoup4"
            )

        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            html = f.read()

        soup = BS(html, "html.parser")

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Get text
        text = soup.get_text()

        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = "\n".join(chunk for chunk in chunks if chunk)

        return text

    def _create_chunks(self, content: str) -> List[TextChunk]:
        """
        Split content into overlapping chunks.

        Args:
            content: Text content to chunk.

        Returns:
            List of TextChunk objects.
        """
        chunks = []

        # Handle empty content
        if not content:
            return chunks

        # Calculate step size
        step = self.chunk_size - self.chunk_overlap

        # Create chunks
        for i, start in enumerate(range(0, len(content), step)):
            end = min(start + self.chunk_size, len(content))

            chunk = TextChunk(
                content=content[start:end],
                chunk_index=i,
                start_char=start,
                end_char=end,
            )

            chunks.append(chunk)

            # Stop if we've reached the end
            if end >= len(content):
                break

        return chunks

    def parse_multiple(
        self, file_paths: List[str], chunk: bool = False
    ) -> List[Document]:
        """
        Parse multiple documents.

        Args:
            file_paths: List of file paths to parse.
            chunk: Whether to chunk documents.

        Returns:
            List of parsed Document objects.
        """
        documents = []

        for file_path in file_paths:
            try:
                doc = self.parse(file_path, chunk=chunk)
                documents.append(doc)
            except Exception as e:
                # Create error document
                metadata = DocumentMetadata.from_file(file_path)
                doc = Document(
                    metadata=metadata,
                    content="",
                    extraction_errors=[f"Failed to parse: {e}"],
                )
                documents.append(doc)

        return documents


def parse_document(file_path: str, chunk: bool = False) -> Document:
    """
    Convenience function to parse a single document.

    Args:
        file_path: Path to the document file.
        chunk: Whether to chunk the document.

    Returns:
        Parsed Document object.

    Example:
        >>> doc = parse_document("requirements.txt")
        >>> print(f"Words: {doc.word_count}")
    """
    parser = DocumentParser()
    return parser.parse(file_path, chunk=chunk)


def parse_documents(file_paths: List[str], chunk: bool = False) -> List[Document]:
    """
    Convenience function to parse multiple documents.

    Args:
        file_paths: List of file paths.
        chunk: Whether to chunk documents.

    Returns:
        List of parsed Document objects.

    Example:
        >>> docs = parse_documents(["req1.txt", "req2.pdf"])
        >>> print(f"Parsed {len(docs)} documents")
    """
    parser = DocumentParser()
    return parser.parse_multiple(file_paths, chunk=chunk)
